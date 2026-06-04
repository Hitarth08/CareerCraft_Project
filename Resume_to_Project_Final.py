# %%
# =========================
# IMPORTS
# =========================
import fitz
from docx import Document
import re
# import spacy
import json
# import tkinter as tk
# from tkinter import filedialog
from groq import Groq
import os
from dotenv import load_dotenv
load_dotenv()

# Load spaCy
# nlp = spacy.load("en_core_web_sm")

# =========================
# GROQ CLIENT SETUP
# =========================
client = Groq(api_key=os.environ.get("GROQ_API_KEY"))  # 🔑 Paste your key here


# =========================
# LLM CALL HELPER
# =========================
def call_llm(prompt, max_tokens=800):
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",   # Same llama3, runs on Groq hardware (very fast)
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=max_tokens
    )
    return response.choices[0].message.content

try:
    import streamlit as st
    api_key = st.secrets["GROQ_API_KEY"]
except:
    from dotenv import load_dotenv
    load_dotenv()
    api_key = os.environ.get("GROQ_API_KEY")

client = Groq(api_key=api_key)

# =========================
# FILE PICKER
# =========================
# def select_file():
#     root = tk.Tk()
#     root.withdraw()
#     root.lift()
#     root.attributes('-topmost', True)

#     file_path = filedialog.askopenfilename(
#         title="Choose Resume File",
#         filetypes=[
#             ("PDF files", "*.pdf"),
#             ("Word files", "*.docx"),
#             ("All files", "*.*")
#         ]
#     )

#     root.destroy()
#     return file_path


# =========================
# EXTRACT TEXT FROM PDF
# =========================
def extract_text_from_pdf(file):
    text = ""
    pdf = fitz.open(stream=file.read(), filetype="pdf")
    for page in pdf:
        text += page.get_text()
    return text


# =========================
# EXTRACT TEXT FROM DOCX
# =========================
def extract_text_from_docx(file):
    doc = Document(file)
    text = [para.text for para in doc.paragraphs]
    return "\n".join(text)


# =========================
# HANDLE FILE TYPES
# =========================
def extract_text(file, filename):
    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(file)
    else:
        raise ValueError("Unsupported format")


# =========================
# CLEAN TEXT
# =========================
def clean_text(text):
    text = re.sub(r'\n+', ' ', text)
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)  # remove markdown links
    return text.strip()


# =========================
# NAME (spaCy fallback)
# =========================
def extract_name_spacy(text):
    # doc = nlp(text[:1000])
    # for ent in doc.ents:
    #     if ent.label_ == "PERSON":
    #         return ent.text
    return "Not Found"


# =========================
# ROBUST JSON EXTRACTOR
# =========================
def extract_json(text):
    start = text.find('{')
    if start == -1:
        return None
    depth = 0
    for i in range(start, len(text)):
        if text[i] == '{':
            depth += 1
        elif text[i] == '}':
            depth -= 1
            if depth == 0:
                return text[start:i+1]
    return None


# =========================
# LLM EXTRACTION
# =========================
def extract_with_llm(text):

    text = text[:1500]

    prompt = f"""
    You are a resume parser.

    Extract ONLY the following fields:

    1. Name
    2. Skills (list all technical + soft skills). Most important — find skills even if scattered. Include all industry-relevant skills.
    3. Education - Degree name and college name. Only undergraduate and masters degrees.
    4. Projects - All project names or competition names only (no descriptions).
    5. Work_Experience_Company - List of company names only (not colleges or addresses).
    6. Work_Experience_in_Years - Total FULL-TIME work experience in years only.
      - Do NOT count internships as experience
      - If only internships exist or no experience, return ""
      - If less than 1 year, return the number of months as "X months" (e.g. "8 months")
      - If 1 year or more, return as a decimal number (e.g. "1.5", "2", "3.5")
    7. Role - Exactly 3 roles the person is most suitable for based on their skills.

    DO NOT extract: Address, Phone, Email, Project descriptions.

    Return STRICT JSON only — no explanation, no markdown, no extra text:

    {{
        "Name": "",
        "Skills": [],
        "Education": [
            {{
                "degree": "",
                "college": ""
            }}
        "Projects": [],
        "Work_Experience_Company": [],
        "Work_Experience_in_Years": "",
        "Role": []
    }}

    Resume Text:
    {text}
    """

    result = call_llm(prompt, max_tokens=800)
    json_str = extract_json(result)

    if json_str:
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            return {"error": f"JSON parsing failed: {str(e)}"}
    else:
        return {"error": "No JSON found", "raw": result[:300]}


# =========================
# COMPANY RECOMMENDATIONS
# =========================
def generate_company_recommendations(profile):

    profile_str = json.dumps(profile)[:800]

    prompt = f"""
    You are an expert Indian job market career advisor with deep knowledge of the Indian hiring landscape.

    Based on the following candidate profile, suggest companies they are MOST LIKELY to get hired at.

    Profile:
    {profile_str}

    Tiering Rules — this is the most important part:

    Tier 1 — Top Stretch Companies:
    - Top MNCs (Google, Microsoft, Amazon, Adobe, Goldman Sachs etc.)
    - Top Indian product companies (Zepto, CRED, Razorpay, Meesho, Groww etc.)
    - Only suggest these if the candidate has: strong skills + good college (NIT/IIT/BITS/VIT/top private) OR 2+ years experience at known companies
    - If the candidate is a fresher from an average college with basic skills → do NOT put MNCs here

    Tier 2 — Good Fit Companies:
    - Mid-size MNCs (Accenture, Capgemini, Cognizant, Mphasis, LTIMindtree etc.)
    - Growing Indian startups (Series A/B funded)
    - Suggest based on domain match — if AI/ML skills → AI startups, if web → product startups
    - These should be realistic for the candidate's actual profile

    Tier 3 — Safe Bets (Very Likely to Hire):
    - Small IT service companies, MSMEs, regional companies
    - For freshers from average colleges → local IT firms, small agencies, BPOs with tech roles
    - For experienced candidates → mid-size service firms in their domain
    - These must be companies the person can almost certainly get into given their current profile

    Key factors to weigh heavily:
    - College tier: IIT/IIM/NIT/BITS = premium tier access | Average private college = start from Tier 3
    - Skills depth: 10+ strong relevant skills vs 3-4 basic skills changes everything
    - Work experience: Even 1 internship at a known company upgrades recommendations
    - Projects: Strong AI/ML/full-stack projects improve tier access significantly
    - Fresher vs experienced: A fresher should never have Google in Tier 1 unless exceptional profile
    - All companies must be real companies actively hiring in India (2026-2027)
    - Do NOT default to IT companies — match companies to the candidate's actual domain
    - For non-IT profiles (Finance, Marketing, HR, Operations, Healthcare and others) 
    suggest companies from THEIR industry, not software companies
    - Only suggest IT/tech companies if the candidate has IT skills or background

    Give exactly 3 companies per tier.
    For each company give a 1-line reason specific to THIS candidate's profile — not generic reasons.

    Return STRICT JSON only — no explanation, no markdown, no extra text:

    {{
        "Company_Recommendations": {{
            "Tier_1_Dream": [
                {{"Company": "", "Reason": ""}}
            ],
            "Tier_2_Good_Fit": [
                {{"Company": "", "Reason": ""}}
            ],
            "Tier_3_Safe_Bet": [
                {{"Company": "", "Reason": ""}}
            ]
        }}
    }}

    Candidate Profile:
    {profile_str}
    """

    result = call_llm(prompt, max_tokens=800)
    json_str = extract_json(result)

    if json_str:
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            return {"error": f"JSON parsing failed: {str(e)}", "raw": result[:300]}
    else:
        return {"error": "No JSON found", "raw": result[:300]}


# =========================
# MAIN PROCESS
# =========================
def process_resume(file_path):

    with open(file_path, "rb") as f:
        raw_text = extract_text(f, file_path)

    cleaned_text = clean_text(raw_text)
    llm_data = extract_with_llm(cleaned_text)

    if llm_data.get("Name") in ["", "Not Found", None]:
        llm_data["Name"] = extract_name_spacy(cleaned_text)

    return llm_data


# =========================
# RUN PROGRAM
# =========================
#if __name__ == "__main__":

    print("Opening file chooser...")
    file_path = select_file()

    if not file_path:
        print("❌ No file selected")
    else:
        print("✅ Selected file:", file_path)

        print("\nProcessing Resume...\n")
        resume_profile = process_resume(file_path)

        print("\n📄 EXTRACTED PROFILE:\n")
        print(json.dumps(resume_profile, indent=4))

        print("\n🏢 GENERATING COMPANY RECOMMENDATIONS...\n")
        companies = generate_company_recommendations(resume_profile)
        print(json.dumps(companies, indent=4))



# =========================
# ATS RESUME ANALYZER
# =========================
def analyze_resume_ats(raw_text, profile):
    
    resume_text = raw_text[:2000]
    profile_str = json.dumps(profile)[:800]

    prompt = f"""
    You are an expert ATS (Applicant Tracking System) specialist and resume coach
    with deep knowledge of how ATS systems work in 2026.

    Analyze the following resume and give detailed ATS optimization recommendations.

    Resume Text:
    {resume_text}

    Extracted Profile:
    {profile_str}

    Analyze and return the following:

    1. ATS_Score (0-100): Overall ATS compatibility score based on:
       - Keyword density and relevance
       - Formatting compatibility
       - Section structure
       - Quantifiable achievements
       - Action verbs usage

    2. Critical_Issues: Things that will cause ATS to REJECT or heavily penalize the resume
       (max 5 issues, only real problems)

    3. Missing_Keywords: Important industry keywords missing from the resume
       based on the person's domain and skills

    4. Section_Feedback: For each section provide:
       - Status: "Strong" / "Needs Improvement" / "Missing"
       - Feedback: What is currently there and what's wrong
       - Suggestion: Exactly what to write or add to improve it

    5. Quick_Wins: Top 5 changes that will immediately improve ATS score

    6. Formatting_Issues: Any formatting problems that confuse ATS parsers

    7. Strong_Points: What the resume is already doing well for ATS

    8. Projected_ATS_Score_After_Changes: Estimated ATS score (0-100)
       if the candidate implements all Quick_Wins and Section_Feedback suggestions

    Return STRICT JSON only — no explanation, no markdown, no HTML tags anywhere in values.
    All string values must be plain text only — no <div>, <span>, or any HTML tags.

    {{
        "ATS_Score": 0,
        "Score_Breakdown": {{
            "Keywords": 0,
            "Formatting": 0,
            "Structure": 0,
            "Achievements": 0,
            "Action_Verbs": 0
        }},
        "Critical_Issues": [],
        "Missing_Keywords": [],
        "Section_Feedback": {{
            "Summary": "",
            "Skills": "",
            "Experience": "",
            "Projects": "",
            "Education": ""
        }},
        "Quick_Wins": [],
        "Formatting_Issues": [],
        "Strong_Points": []
        "Projected_ATS_Score_After_Changes": 0
    }}

    Resume Text:
    {resume_text}
    """

    result = call_llm(prompt, max_tokens=2000)
    json_str = extract_json(result)

    if json_str:
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            return {"error": f"JSON parsing failed: {str(e)}"}
    return {"error": "No JSON found"}


# =========================
# PROJECT RECOMMENDATIONS
# =========================

def get_user_goal():
    print("\n🎯 WHAT IS YOUR GOAL?")
    print("1. Upskilling (improve my current profile)")
    print("2. Job (target specific companies and roles)")
    
    while True:
        choice = input("\nEnter 1 or 2: ").strip()
        if choice in ["1", "2"]:
            return "upskilling" if choice == "1" else "job"
        print("❌ Invalid input. Please enter 1 or 2.")


def get_job_targets():
    print("\n🏢 ENTER YOUR TARGET COMPANIES (3 companies)")
    companies = []
    for i in range(1, 4):
        company = input(f"  Company {i}: ").strip()
        companies.append(company)

    print("\n💼 ENTER YOUR TARGET ROLE")
    role = input("  Role (e.g. Data Scientist, ML Engineer): ").strip()

    return companies, role

def generate_project_recommendations(profile):

    goal = get_user_goal()
    profile_str = json.dumps(profile)[:1000]

    # ─── UPSKILLING PATH ───────────────────────────────────────────
    if goal == "upskilling":

        prompt = f"""
        You are an expert career mentor and project advisor.
        - If the candidate is from a non-IT domain (Finance, Marketing, HR, Operations, 
        Healthcare, Manufacturing etc.) suggest projects relevant to THEIR domain
        - Do NOT default to software/ML projects for non-IT profiles
        The candidate wants to UPSKILL and improve their current profile.

        Based on the following profile, suggest EXACTLY 5 project ideas that will
        help them grow beyond what they have already done.

        Profile:
        {profile_str}

        Requirements:
        - Do NOT suggest projects similar to existing ones in the profile
        - Projects must fill skill gaps or extend current strengths
        - Must be INDUSTRY-RELEVANT (2026 level)
        - Adjust difficulty based on profile:
            - Beginner: for freshers or limited projects
            - Intermediate: some experience or projects
            - Expert: strong profile with work experience
        - Distribute levels smartly (e.g. 2 Beginner, 2 Intermediate, 1 Expert)

        Each project must include:
        - Title
        - Level (Beginner / Intermediate / Expert)
        - Description (5-6 lines explaining what the project does)
        - Why_Suitable (why this helps the candidate upskill based on their profile)
        - Tech_Stack (list of tools/technologies)

        Return STRICT JSON only — no explanation, no markdown:

        {{
            "Goal": "Upskilling",
            "Recommended_Projects": [
                {{
                    "Title": "",
                    "Level": "",
                    "Description": "",
                    "Why_Suitable": "",
                    "Tech_Stack": []
                }}
            ]
        }}
        """

        result = call_llm(prompt, max_tokens=2000)

    # ─── JOB PATH ──────────────────────────────────────────────────
    else:

        target_companies, target_role = get_job_targets()
        companies_str = ", ".join(target_companies)

        prompt = f"""
        You are an expert career mentor and project advisor.

        The candidate is targeting a JOB and wants project recommendations
        that will maximize their chances of getting hired.

        Candidate Profile:
        {profile_str}

        Target Companies: {companies_str}
        Target Role: {target_role}

        Requirements:
        - Suggest EXACTLY 5 projects that directly align with what these companies
          look for in a {target_role}
        - Research what tech stacks, domains, and problem types these companies work on
          and base the projects on that
        - Do NOT suggest projects similar to existing ones in the profile
        - Projects should demonstrate skills that are specifically valued for
          {target_role} at {companies_str}
        - Must be INDUSTRY-RELEVANT (2026 level)
        - Adjust difficulty based on profile:
            - Beginner: for freshers or limited projects
            - Intermediate: some experience or projects
            - Expert: strong profile with work experience
        - Distribute levels smartly (e.g. 2 Beginner, 2 Intermediate, 1 Expert)

        Each project must include:
        - Title
        - Level (Beginner / Intermediate / Expert)
        - Description (5-6 lines explaining what the project does)
        - Why_Suitable (why this project helps get hired at the target companies
          for the target role — be specific)
        - Tech_Stack (list of tools/technologies)

        Return STRICT JSON only — no explanation, no markdown:

        {{
            "Goal": "Job",
            "Target_Companies": [],
            "Target_Role": "",
            "Recommended_Projects": [
                {{
                    "Title": "",
                    "Level": "",
                    "Description": "",
                    "Why_Suitable": "",
                    "Tech_Stack": []
                }}
            ]
        }}
        """

        result = call_llm(prompt, max_tokens=2000)



    # ─── PARSE AND RETURN ──────────────────────────────────────────
    json_str = extract_json(result)

    if json_str:
        try:
            return json.loads(json_str)
        except json.JSONDecodeError as e:
            return {"error": f"JSON parsing failed: {str(e)}", "raw": result[:300]}
    else:
        return {"error": "No JSON found", "raw": result[:300]}
    
# if __name__ == "__main__":

#     print("Opening file chooser...")
#     file_path = select_file()

#     if not file_path:
#         print("❌ No file selected")
#     else:
#         print("✅ Selected file:", file_path)

#         print("\nProcessing Resume...\n")
#         resume_profile = process_resume(file_path)

#         print("\n📄 EXTRACTED PROFILE:\n")
#         print(json.dumps(resume_profile, indent=4))

#         print("\n🏢 GENERATING COMPANY RECOMMENDATIONS...\n")
#         #companies = generate_company_recommendations(resume_profile)
#         #print(json.dumps(companies, indent=4))
# ── ATS ANALYSIS ──────────────────────────────────────
        print("\n🎯 GENERATING ATS ANALYSIS...\n")

        with open(file_path, "rb") as f:
            from docx import Document
            import fitz

            if file_path.endswith(".pdf"):
                pdf = fitz.open(stream=f.read(), filetype="pdf")
                raw_text = "".join(page.get_text() for page in pdf)
            else:
                doc = Document(f)
                raw_text = "\n".join([p.text for p in doc.paragraphs])

        cleaned = clean_text(raw_text)
        ats_result = analyze_resume_ats(cleaned, resume_profile)
        print(json.dumps(ats_result, indent=4))

        print("\n🚀 GENERATING PROJECT RECOMMENDATIONS...\n")
        #projects = generate_project_recommendations(resume_profile)
        #print(json.dumps(projects, indent=4))

# %%



